"""
ui/richtext_widget.py — Editor Rich Text (Jodit + QWebEngineView)
NotePadPQ

Editor WYSIWYG completo basato su Jodit (https://jodit-editor.com),
incorporato in QWebEngineView con comunicazione bidirezionale via QWebChannel.

Formati supportati:
  Lettura : .docx (mammoth), .odt/.rtf (pandoc), .html/.htm (nativo)
  Scrittura: .html (nativo), .docx (htmldocx/python-docx), .pdf (Qt print),
             .odt/.rtf (pandoc)

Dipendenze opzionali (degradano gracefully):
  PyQt6-WebEngine  — obbligatoria per il widget
  mammoth          — lettura DOCX
  htmldocx         — scrittura DOCX
  python-docx      — fallback scrittura DOCX
  pypandoc         — ODT/RTF conversione
"""

from __future__ import annotations

import json
import tempfile
import shutil
import urllib.request
from pathlib import Path
from typing import Optional, TYPE_CHECKING

from PyQt6.QtCore import Qt, QObject, QThread, QTimer, QEventLoop, QUrl, pyqtSignal, pyqtSlot
from PyQt6.QtGui import QKeySequence, QShortcut
from PyQt6.QtWidgets import (
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QToolBar,
    QToolButton,
    QComboBox,
    QLabel,
    QFileDialog,
    QMessageBox,
    QProgressDialog,
    QSizePolicy,
)

try:
    from PyQt6.QtWebEngineWidgets import QWebEngineView
    from PyQt6.QtWebEngineCore import (
        QWebEngineSettings,
        QWebEngineScript,
        QWebEnginePage,
    )
    from PyQt6.QtWebChannel import QWebChannel

    WEBENGINE_OK = True
except ImportError:
    WEBENGINE_OK = False

from i18n.i18n import tr


def _load_rt_icon(svg_file: str, widget: QWidget, size: int = 18) -> "QIcon":
    """Carica un'icona SVG dal set attivo, sostituendo currentColor con il colore
    testo corrente della palette. Ritorna QIcon vuota se non trovata."""
    from PyQt6.QtGui import QIcon, QPalette, QPixmap
    from config.settings import Settings

    icon_path = Path(__file__).parent.parent / "icons" / "lucide" / svg_file

    if not icon_path.exists():
        return QIcon()

    color = widget.palette().color(QPalette.ColorRole.WindowText).name()
    try:
        svg_data = icon_path.read_bytes().replace(b"currentColor", color.encode())
        pm = QPixmap()
        if pm.loadFromData(svg_data, "SVG") and not pm.isNull():
            return QIcon(pm)
    except Exception:
        pass
    return QIcon()


# ── Asset Jodit ───────────────────────────────────────────────────────────────

_ASSETS_DIR = Path(__file__).parent / "assets" / "jodit"

# Jodit 4 — MIT licence, no dependencies, single-file distribution
# URL senza versione: jsdelivr risolve sempre all'ultima stabile
_JODIT_JS_URL = "https://cdn.jsdelivr.net/npm/jodit/es2021/jodit.min.js"
_JODIT_CSS_URL = "https://cdn.jsdelivr.net/npm/jodit/es2021/jodit.min.css"

_JODIT_JS = _ASSETS_DIR / "jodit.min.js"
_JODIT_CSS = _ASSETS_DIR / "jodit.min.css"
_ACTIVE_THREADS: set[QThread] = set()


def _keep_thread_alive(thread: QThread) -> None:
    _ACTIVE_THREADS.add(thread)
    thread.finished.connect(lambda thread=thread: _ACTIVE_THREADS.discard(thread))


class _JoditDownloadWorker(QThread):
    progress = pyqtSignal(int, str)
    completed = pyqtSignal(bool, str)

    def __init__(self):
        super().__init__()
        self._cancelled = False

    def cancel(self) -> None:
        self._cancelled = True

    def run(self) -> None:
        files = [(_JODIT_JS_URL, _JODIT_JS), (_JODIT_CSS_URL, _JODIT_CSS)]
        try:
            _ASSETS_DIR.mkdir(parents=True, exist_ok=True)
            for i, (url, dest) in enumerate(files):
                if self._cancelled:
                    raise InterruptedError
                label = f"Download {dest.name}..."
                self.progress.emit(i * 50, label)
                req = urllib.request.Request(
                    url, headers={"User-Agent": "NotePadPQ/1.0 (richtext downloader)"}
                )
                part = dest.with_suffix(dest.suffix + ".part")
                with urllib.request.urlopen(req, timeout=30) as resp, open(part, "wb") as f:
                    total = int(resp.headers.get("Content-Length", 0))
                    downloaded = 0
                    while not self._cancelled:
                        chunk = resp.read(32768)
                        if not chunk:
                            break
                        f.write(chunk)
                        downloaded += len(chunk)
                        if total:
                            self.progress.emit(
                                i * 50 + int(downloaded / total * 50), label
                            )
                if self._cancelled:
                    part.unlink(missing_ok=True)
                    raise InterruptedError
                part.replace(dest)
            self.completed.emit(True, "")
        except InterruptedError:
            for _, dest in files:
                dest.with_suffix(dest.suffix + ".part").unlink(missing_ok=True)
            self.completed.emit(False, "")
        except Exception as exc:
            for _, dest in files:
                dest.with_suffix(dest.suffix + ".part").unlink(missing_ok=True)
            self.completed.emit(False, str(exc))


class JoditDownload(QObject):
    """Gestisce la UI del download, eseguito interamente nel worker."""

    finished = pyqtSignal(bool)

    def __init__(self, parent: Optional[QWidget] = None):
        super().__init__(parent)
        self._ok = False
        self._error = ""
        self._cancelled = False
        self._worker = _JoditDownloadWorker()
        self._dialog = QProgressDialog(
            tr("richtext.downloading_jodit"), tr("button.cancel"), 0, 100, parent
        )
        self._dialog.setWindowTitle("NotePadPQ")
        self._dialog.setWindowModality(Qt.WindowModality.ApplicationModal)
        self._dialog.setMinimumDuration(0)
        self._dialog.canceled.connect(self._worker.cancel)
        self._worker.progress.connect(self._on_progress)
        self._worker.completed.connect(self._on_completed)
        self._worker.finished.connect(self._on_thread_finished)
        self._worker.finished.connect(self._worker.deleteLater)

    def start(self) -> None:
        self._dialog.show()
        _keep_thread_alive(self._worker)
        self._worker.start()

    def cancel(self) -> None:
        """Annulla il download senza mostrare errori per la cancellazione."""
        self._cancelled = True
        self._worker.cancel()
        self._dialog.close()

    @pyqtSlot(int, str)
    def _on_progress(self, value: int, label: str) -> None:
        self._dialog.setLabelText(label)
        self._dialog.setValue(value)

    @pyqtSlot(bool, str)
    def _on_completed(self, ok: bool, error: str) -> None:
        self._ok = ok and not self._cancelled
        self._error = "" if self._cancelled else error

    @pyqtSlot()
    def _on_thread_finished(self) -> None:
        self._dialog.close()
        if self._error:
            QMessageBox.critical(
                self.parent(), "NotePadPQ",
                tr("richtext.download_failed", exc=self._error),
            )
        self.finished.emit(self._ok)


def download_jodit(parent: Optional[QWidget] = None) -> JoditDownload:
    controller = JoditDownload(parent)
    controller.start()
    return controller


def _download_jodit(parent: Optional[QWidget] = None) -> bool:
    """API storica: attende l'esito mantenendo reattivo l'event loop Qt."""
    result = [False]
    loop = QEventLoop()
    controller = download_jodit(parent)

    def done(ok: bool) -> None:
        result[0] = ok
        loop.quit()

    controller.finished.connect(done)
    loop.exec()
    return result[0]


def ensure_jodit(parent: Optional[QWidget] = None) -> bool:
    """Verifica/scarica gli asset Jodit. Ritorna True se pronti."""
    if _JODIT_JS.exists() and _JODIT_CSS.exists():
        return True
    return _download_jodit(parent)


# ── Conversione formati ───────────────────────────────────────────────────────


class RichTextIO:
    """Conversione bidirezionale tra HTML interno e formati file."""

    @staticmethod
    def load_html(path: Path) -> tuple[str, str]:
        """Legge un file e ritorna (html_content, error_message)."""
        ext = path.suffix.lower()

        if ext in (".html", ".htm"):
            try:
                return path.read_text(encoding="utf-8", errors="replace"), ""
            except Exception as e:
                return "", str(e)

        if ext == ".docx":
            return RichTextIO._docx_to_html(path)

        if ext == ".doc":
            return RichTextIO._doc_to_html(path)

        if ext in (".odt", ".rtf"):
            return RichTextIO._pandoc_to_html(path)

        try:
            return path.read_text(encoding="utf-8", errors="replace"), ""
        except Exception as e:
            return "", str(e)

    @staticmethod
    def _docx_to_html(path: Path) -> tuple[str, str]:
        try:
            import mammoth

            with open(path, "rb") as f:
                result = mammoth.convert_to_html(f)
            html = result.value
            if result.messages:
                # Log warnings but don't fail
                pass
            return html, ""
        except ImportError:
            return RichTextIO._pandoc_to_html(path)
        except Exception as e:
            return "", str(e)

    @staticmethod
    def _doc_to_html(path: Path) -> tuple[str, str]:
        """Converte .doc (Word 97-2003) → HTML via LibreOffice → mammoth."""
        import subprocess, tempfile, shutil

        tmp = Path(tempfile.mkdtemp(prefix="npq_doc_"))
        try:
            res = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(tmp),
                    str(path),
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if res.returncode != 0:
                return "", f"LibreOffice: {res.stderr.strip() or res.stdout.strip()}"
            docx = tmp / (path.stem + ".docx")
            if not docx.exists():
                # LibreOffice may keep original stem with spaces/accents
                candidates = list(tmp.glob("*.docx"))
                if not candidates:
                    return "", "LibreOffice non ha prodotto un file .docx"
                docx = candidates[0]
            return RichTextIO._docx_to_html(docx)
        except FileNotFoundError:
            return (
                "",
                "LibreOffice non trovato — installa libreoffice per aprire file .doc",
            )
        except Exception as e:
            return "", str(e)
        finally:
            shutil.rmtree(tmp, ignore_errors=True)

    @staticmethod
    def _extract_body(html: str) -> str:
        """Estrae il contenuto del <body> da un documento HTML completo."""
        import re

        m = re.search(r"<body(?:[^>]*)>(.*?)</body>", html, re.DOTALL | re.IGNORECASE)
        return m.group(1).strip() if m else html

    @staticmethod
    def _pandoc_to_html(path: Path) -> tuple[str, str]:
        fmt = path.suffix.lstrip(".")
        try:
            import pypandoc

            # Prova --embed-resources (pandoc 3+), fallback a --self-contained
            for embed in ("--embed-resources", "--self-contained"):
                try:
                    html = pypandoc.convert_file(
                        str(path), "html", extra_args=["--standalone", embed]
                    )
                    return RichTextIO._extract_body(html), ""
                except Exception:
                    continue
        except ImportError:
            pass
        # Fallback: pandoc subprocess con immagini embedded come base64
        import subprocess

        try:
            for embed in ("--embed-resources", "--self-contained"):
                result = subprocess.run(
                    [
                        "pandoc",
                        "-f",
                        fmt,
                        "-t",
                        "html",
                        "--standalone",
                        embed,
                        str(path),
                    ],
                    capture_output=True,
                    text=True,
                    timeout=60,
                )
                if result.returncode == 0:
                    return RichTextIO._extract_body(result.stdout), ""
            return "", result.stderr
        except FileNotFoundError:
            return "", tr("richtext.pandoc_missing")
        except Exception as e:
            return "", str(e)

    @staticmethod
    def save_html(html: str, path: Path) -> str:
        """Salva HTML nel formato del file. Ritorna stringa errore o ''."""
        ext = path.suffix.lower()

        if ext in (".html", ".htm"):
            try:
                path.write_text(RichTextIO._full_html(html), encoding="utf-8")
                return ""
            except Exception as e:
                return str(e)

        if ext == ".docx":
            return RichTextIO._html_to_docx(html, path)

        if ext in (".odt", ".rtf"):
            return RichTextIO._html_to_pandoc(html, path)

        try:
            path.write_text(html, encoding="utf-8")
            return ""
        except Exception as e:
            return str(e)

    @staticmethod
    def _html_to_docx(html: str, path: Path) -> str:
        # Preferisce htmldocx, fallback python-docx + BeautifulSoup
        try:
            from htmldocx import HtmlToDocx

            parser = HtmlToDocx()
            doc = parser.parse_html_string(html)
            doc.save(str(path))
            return ""
        except ImportError:
            pass
        try:
            from docx import Document

            doc = Document()
            # Inserisce il testo come plain text (degraded)
            from html.parser import HTMLParser

            class _TextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text = []

                def handle_data(self, data):
                    self.text.append(data)

            ex = _TextExtractor()
            ex.feed(html)
            for chunk in " ".join(ex.text).split("\n"):
                if chunk.strip():
                    doc.add_paragraph(chunk.strip())
            doc.save(str(path))
            return ""
        except ImportError:
            return tr("richtext.docx_lib_missing")
        except Exception as e:
            return str(e)

    @staticmethod
    def _html_to_pandoc(html: str, path: Path) -> str:
        import subprocess

        with tempfile.NamedTemporaryFile(
            suffix=".html", delete=False, mode="w", encoding="utf-8"
        ) as f:
            f.write(RichTextIO._full_html(html))
            tmp = f.name
        try:
            fmt = path.suffix.lstrip(".")
            result = subprocess.run(
                ["pandoc", "-f", "html", "-t", fmt, tmp, "-o", str(path)],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode != 0:
                return result.stderr
            return ""
        except FileNotFoundError:
            return tr("richtext.pandoc_missing")
        except Exception as e:
            return str(e)
        finally:
            Path(tmp).unlink(missing_ok=True)

    @staticmethod
    def _full_html(body: str) -> str:
        return (
            "<!DOCTYPE html><html><head>"
            '<meta charset="utf-8">'
            "<style>body{font-family:serif;font-size:12pt;line-height:1.5;"
            "max-width:860px;margin:40px auto;padding:0 20px}</style>"
            f"</head><body>{body}</body></html>"
        )


class _RichTextIOWorker(QThread):
    """Esegue lettura e conversione senza accedere a oggetti UI."""

    completed = pyqtSignal(str, str, bool)  # result, error, cancelled

    def __init__(self, operation: str, path: Path, html: str = ""):
        super().__init__()
        self._operation = operation
        self._path = path
        self._html = html
        self._cancelled = False

    def cancel(self) -> None:
        self._cancelled = True

    def run(self) -> None:
        if self._cancelled:
            self.completed.emit("", "", True)
            return
        try:
            if self._operation == "load":
                result, error = RichTextIO.load_html(self._path)
            else:
                result, error = "", RichTextIO.save_html(self._html, self._path)
        except Exception as exc:
            result, error = "", str(exc)
        self.completed.emit(result, error, self._cancelled)


# ── QWebChannel bridge ────────────────────────────────────────────────────────


class _EditorBridge(QObject):
    """Oggetto esposto a JavaScript via QWebChannel."""

    content_changed = pyqtSignal()
    js_ready = pyqtSignal()

    def __init__(self, widget: "RichTextWidget"):
        super().__init__()
        self._widget = widget

    @pyqtSlot()
    def ready(self):
        """Chiamato da JS quando Jodit è inizializzato."""
        self.js_ready.emit()

    @pyqtSlot()
    def contentChanged(self):
        """Chiamato da JS ad ogni modifica dell'utente."""
        self._widget._set_modified(True)
        self.content_changed.emit()

    @pyqtSlot()
    def printRequested(self):
        """Chiamato da JS quando Jodit invoca window.print() o il pulsante Print."""
        self._widget._print_document()

    @pyqtSlot(str)
    def log(self, msg: str):
        pass  # debug hook


# ── Template HTML ─────────────────────────────────────────────────────────────

_HTML_TEMPLATE = """\
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<link rel="stylesheet" href="jodit.min.css">
<style>
  html, body {{
    margin: 0; padding: 0;
    width: 100%; height: 100%;
    overflow: hidden;
  }}
  #editor {{ display: block; width: 100%; }}
  .jodit-container {{
    width: 100% !important;
    max-width: none !important;
    box-sizing: border-box !important;
  }}
  .jodit-workplace,
  .jodit-wysiwyg,
  .jodit-workplace__area,
  .jodit-wysiwyg-iframe {{
    width: 100% !important;
    max-width: none !important;
  }}
</style>
</head>
<body>
<textarea id="editor"></textarea>
<script src="qwebchannel.js"></script>
<script src="jodit.min.js"></script>
<script>
"use strict";

var bridge  = null;
var jodit   = null;
var _pendingContent = null;
var _changing = false;

// ── Resize helper ─────────────────────────────────────────────────────────────
window.joditForceSize = function() {{
  if (!jodit || !jodit.container) return;
  var h = window.innerHeight;
  jodit.container.style.height = h + "px";
  jodit.e.fire("resize");
}};

window.addEventListener("resize", function() {{
  window.joditForceSize();
}});

// ── Jodit config ──────────────────────────────────────────────────────────────
var JODIT_CONFIG = {{
  height: window.innerHeight,
  width:  '100%',
  minHeight: 400,
  autofocus: false,
  spellcheck: true,
  language: "{lang}",
  toolbarAdaptive: false,
  toolbarSticky: true,
  showCharsCounter: false,
  showWordsCounter: true,
  showXPathInStatusbar: false,
  allowResizeY: false,
  iframe: false,
  useAceEditor: false,
  uploader: {{ insertImageAsBase64URI: true }},
  buttons: [
    "bold","italic","underline","strikethrough","eraser","|",
    "superscript","subscript","|",
    "ul","ol","|",
    "outdent","indent","|",
    "font","fontsize","paragraph","brush","|",
    "align","|",
    "table","image","link","|",
    "undo","redo","|",
    "hr","copyformat","|",
    "find","|",
    "source","fullsize"
  ],
  extraButtons: [],
  controls: {{
    paragraph: {{
      list: {{
        p:  "Normale",
        h1: "Titolo 1",
        h2: "Titolo 2",
        h3: "Titolo 3",
        h4: "Titolo 4",
        blockquote: "Citazione",
        pre: "Codice"
      }}
    }}
  }}
}};

// ── Funzioni esposte a Python ─────────────────────────────────────────────────
window.setContent = function(html) {{
  if (jodit) {{
    jodit.value = html;
    jodit.history.clear();
  }} else {{
    _pendingContent = html;
  }}
}};

window.getContent = function() {{
  return jodit ? jodit.value : "";
}};

window.setReadOnly = function(ro) {{
  if (jodit) jodit.setReadOnly(ro);
}};

window.execCommand = function(cmd, val) {{
  if (jodit) jodit.execCommand(cmd, false, val || "");
}};

// ── Inizializzazione ──────────────────────────────────────────────────────────
new QWebChannel(qt.webChannelTransport, function(channel) {{
  bridge = channel.objects.bridge;

  // Fallback: se qualche percorso chiama window.print() direttamente.
  window.print = function() {{ if (bridge) bridge.printRequested(); }};

  jodit = Jodit.make("#editor", JODIT_CONFIG);

  jodit.events.on("afterInit", function() {{
    window.joditForceSize();
    setTimeout(window.joditForceSize, 150);
  }});

  if (_pendingContent !== null) {{
    jodit.value = _pendingContent;
    jodit.history.clear();
    _pendingContent = null;
  }}

  jodit.events.on("change", function() {{
    if (!_changing && bridge) bridge.contentChanged();
  }});

  bridge.ready();
}});
</script>
</body>
</html>
"""


# ── Widget principale ─────────────────────────────────────────────────────────

_SAVE_FILTER = (
    "HTML (*.html *.htm);;"
    "Documento Word (*.docx);;"
    "PDF (*.pdf);;"
    "OpenDocument Text (*.odt);;"
    "RTF (*.rtf);;"
    "Tutti i file (*)"
)


class RichTextWidget(QWidget):
    """
    Tab widget per l'editor rich text WYSIWYG.
    Implementa l'interfaccia attesa da TabManager.add_spreadsheet_tab:
      - file_path      : Path | None
      - modified_changed signal
      - is_modified()
      - save()
    """

    modified_changed = pyqtSignal(bool)
    convert_to_text = pyqtSignal(str, str)  # (content, suggested_name)
    load_finished = pyqtSignal(bool)
    save_finished = pyqtSignal(bool, object)  # success, Path

    def __init__(self, path: Optional[Path] = None, parent: Optional[QWidget] = None):
        super().__init__(parent)
        self.file_path: Optional[Path] = path
        self._modified = False
        self._js_ready = False
        self._pending_html: Optional[str] = None
        self._tmpdir: Optional[Path] = None
        self._io_worker: Optional[_RichTextIOWorker] = None
        self._progress_dlg: Optional[QProgressDialog] = None
        self._io_operation = ""
        self._io_path: Optional[Path] = None
        self._io_save_as = False
        self._io_pending = False
        self._io_ok = False
        self._save_change_serial = 0
        self._change_serial = 0
        self._closed = False

        self._build_ui()
        self._load_jodit_page()

    # ── Costruzione UI ────────────────────────────────────────────────────────

    def _build_ui(self) -> None:
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Toolbar superiore
        self._toolbar = self._make_toolbar()
        layout.addWidget(self._toolbar)

        if not WEBENGINE_OK:
            lbl = QLabel(
                "⚠ PyQt6-WebEngine non installato.\n" "pip install PyQt6-WebEngine"
            )
            lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
            layout.addWidget(lbl)
            return

        # WebEngineView
        from core.webengine import safe_webview

        self._view = safe_webview(self)
        if self._view is None:
            from i18n.i18n import tr as _tr

            lbl = QLabel(_tr("plugin.terminal.no_gl"), self)
            lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
            lbl.setWordWrap(True)
            layout.addWidget(lbl)
            return
        self._view.setSizePolicy(
            QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding
        )
        # Abilita Javascript e accesso file locali
        settings = self._view.settings()
        settings.setAttribute(QWebEngineSettings.WebAttribute.JavascriptEnabled, True)
        settings.setAttribute(
            QWebEngineSettings.WebAttribute.LocalContentCanAccessFileUrls, True
        )
        settings.setAttribute(
            QWebEngineSettings.WebAttribute.LocalContentCanAccessRemoteUrls, False
        )
        layout.addWidget(self._view)

        # Status bar
        self._statusbar = self._make_statusbar()
        layout.addWidget(self._statusbar)

        # Shortcut salvataggio
        QShortcut(QKeySequence("Ctrl+S"), self, self.save)
        QShortcut(QKeySequence("Ctrl+Shift+S"), self, self.save_as)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        if self._js_ready and getattr(self, "_view", None) is not None:
            self._view.page().runJavaScript(
                "if(typeof joditForceSize==='function') joditForceSize();"
            )

    def _make_toolbar(self) -> QWidget:
        from PyQt6.QtCore import QSize
        from PyQt6.QtGui import QIcon

        bar = QToolBar(self)
        bar.setMovable(False)
        bar.setFloatable(False)
        _icon_size = QSize(18, 18)
        bar.setIconSize(_icon_size)

        _save_as_svg = "save.svg"

        def btn(svg_file: str, fallback_text: str, tip: str, slot) -> QToolButton:
            b = QToolButton(bar)
            b.setToolTip(tip)
            b.setIconSize(_icon_size)
            icon = _load_rt_icon(svg_file, self)
            if not icon.isNull():
                b.setIcon(icon)
                b.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonIconOnly)
            else:
                b.setText(fallback_text)
                b.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextOnly)
            b.clicked.connect(slot)
            bar.addWidget(b)
            return b

        btn("save.svg", "💾", tr("action.save") + "  Ctrl+S", self.save)
        btn(_save_as_svg, "📁", tr("richtext.save_as") + "  Ctrl+Shift+S", self.save_as)
        bar.addSeparator()
        btn("printer.svg", "🖨", tr("richtext.print"), lambda _: self._print_document())
        _b_pdf = QToolButton(bar)
        _b_pdf.setText("PDF")
        _b_pdf.setToolTip(tr("richtext.export_pdf"))
        _b_pdf.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextOnly)
        _b_pdf.clicked.connect(lambda _: self._export_pdf())
        bar.addWidget(_b_pdf)
        bar.addSeparator()
        btn("file-text.svg", "✎", tr("richtext.open_as_text"), self._open_as_text)

        # Word count label a destra
        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        bar.addWidget(spacer)
        self._wc_label = QLabel("", bar)
        self._wc_label.setStyleSheet("padding: 0 8px; color: gray;")
        bar.addWidget(self._wc_label)

        return bar

    def _make_statusbar(self) -> QWidget:
        w = QWidget(self)
        w.setFixedHeight(22)
        w.setStyleSheet(
            "QWidget { background: palette(window); border-top: 1px solid palette(mid); }"
            "QLabel  { padding: 0 6px; color: gray; font-size: 11px; }"
        )
        h = QHBoxLayout(w)
        h.setContentsMargins(0, 0, 0, 0)
        h.setSpacing(0)
        self._status_lbl = QLabel("", w)
        h.addWidget(self._status_lbl)
        h.addStretch()
        self._mod_lbl = QLabel("", w)
        h.addWidget(self._mod_lbl)
        return w

    # ── Caricamento pagina Jodit ──────────────────────────────────────────────

    def _load_jodit_page(self) -> None:
        if not WEBENGINE_OK or self._view is None:
            return

        # Crea tmpdir con gli asset di Jodit
        self._tmpdir = Path(tempfile.mkdtemp(prefix="npq_rt_"))
        js_src = _JODIT_JS
        css_src = _JODIT_CSS

        if not (js_src.exists() and css_src.exists()):
            self._status_lbl.setText(tr("richtext.jodit_not_ready"))
            return

        shutil.copy(js_src, self._tmpdir / "jodit.min.js")
        shutil.copy(css_src, self._tmpdir / "jodit.min.css")

        # Ottieni qwebchannel.js dal sistema Qt
        qwc_js = self._get_qwebchannel_js()
        (self._tmpdir / "qwebchannel.js").write_text(qwc_js, encoding="utf-8")

        # Genera la pagina HTML
        from i18n.i18n import I18n

        lang = (
            I18n.instance().current_language()
            if hasattr(I18n.instance(), "current_language")
            else "it"
        )
        jodit_lang = {"it": "it", "en": "en", "de": "de", "fr": "fr", "es": "es"}.get(
            lang, "en"
        )

        html_content = _HTML_TEMPLATE.format(lang=jodit_lang)
        html_path = self._tmpdir / "editor.html"
        html_path.write_text(html_content, encoding="utf-8")

        # Configura QWebChannel
        self._bridge = _EditorBridge(self)
        self._channel = QWebChannel(self._view.page())
        self._channel.registerObject("bridge", self._bridge)
        self._view.page().setWebChannel(self._channel)

        self._bridge.js_ready.connect(self._on_js_ready)
        self._bridge.content_changed.connect(self._on_content_changed)

        self._view.setUrl(QUrl.fromLocalFile(str(html_path)))

    def _get_qwebchannel_js(self) -> str:
        from PyQt6.QtCore import QFile, QIODevice

        f = QFile(":/qtwebchannel/qwebchannel.js")
        if f.open(QIODevice.OpenModeFlag.ReadOnly):
            data = bytes(f.readAll()).decode("utf-8")
            f.close()
            return data
        return ""

    # ── Slot ──────────────────────────────────────────────────────────────────

    def _on_js_ready(self) -> None:
        self._js_ready = True
        if self._pending_html is not None:
            self._js_set_content(self._pending_html)
            self._pending_html = None
        self._update_status()

    def _on_content_changed(self) -> None:
        self._change_serial += 1
        # Aggiorna word count ogni 2s per non appesantire
        if not hasattr(self, "_wc_timer"):
            self._wc_timer = QTimer(self)
            self._wc_timer.setSingleShot(True)
            self._wc_timer.timeout.connect(self._refresh_word_count)
        self._wc_timer.start(2000)

    # ── API pubblica ──────────────────────────────────────────────────────────

    def load_document(self, path: Path) -> bool:
        """Avvia il caricamento. False indica che un'altra operazione è attiva."""
        return self._start_io("load", path)

    def set_html(self, html: str) -> None:
        """Imposta il contenuto HTML nell'editor."""
        if self._js_ready:
            self._js_set_content(html)
        else:
            self._pending_html = html

    def get_html(self) -> str:
        """Legge il contenuto HTML corrente dall'editor (sincrono)."""
        if not WEBENGINE_OK or not self._js_ready:
            return ""
        result: list[Optional[str]] = [None]
        loop = QEventLoop()

        def cb(val: str):
            result[0] = val
            loop.quit()

        self._view.page().runJavaScript("window.getContent()", cb)
        loop.exec()
        return result[0] or ""

    def is_modified(self) -> bool:
        return self._modified

    def is_save_in_progress(self) -> bool:
        """Indica se il contenuto viene acquisito o scritto su disco."""
        return self._io_pending or (
            self._io_worker is not None and self._io_operation == "save"
        )

    def save(self) -> bool:
        """Avvia il salvataggio nel formato originale."""
        if self.file_path is None:
            return self.save_as()
        return self._request_save(self.file_path, False)

    def save_as(self) -> bool:
        import re as _re

        # Mostra il nome senza estensione nel campo del dialog
        if self.file_path:
            default = str(self.file_path.with_suffix(""))
        else:
            default = str(Path.home() / "documento")
        path, selected_filter = QFileDialog.getSaveFileName(
            self, tr("richtext.save_as"), default, _SAVE_FILTER
        )
        if not path:
            return False
        p = Path(path)
        # Se il path non ha estensione (o non corrisponde al filtro scelto),
        # applica l'estensione del filtro selezionato
        m = _re.search(r"\*(\.\w+)", selected_filter)
        if m:
            ext = m.group(1).lower()
            if p.suffix.lower() != ext:
                p = p.with_suffix(ext)
        if p.suffix.lower() == ".pdf":
            ok = self._export_pdf(p)
            self.save_finished.emit(ok, p)
            return ok
        return self._request_save(p, True)

    def print_document(self) -> None:
        """Metodo pubblico: apre la dialog di stampa di sistema (File > Stampa)."""
        self._print_document()

    def _print_document(self) -> None:
        """Apre QPrintDialog e stampa il contenuto del documento."""
        from PyQt6.QtPrintSupport import QPrinter, QPrintDialog
        from PyQt6.QtGui import QTextDocument

        if not WEBENGINE_OK or not self._js_ready:
            return
        # Mostra la dialog prima (sincrona), poi legge l'HTML via callback asincrono
        # per evitare event loop annidati che causano deadlock con QWebChannel.
        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        dlg = QPrintDialog(printer, self.window())
        if dlg.exec() != QPrintDialog.DialogCode.Accepted:
            return
        _printer = printer

        def _do_print(html: str) -> None:
            doc = QTextDocument()
            doc.setHtml(RichTextIO._full_html(html or ""))
            doc.print(_printer)

        self._view.page().runJavaScript("window.getContent()", _do_print)

    def export_pdf(self) -> None:
        """Metodo pubblico per esportare come PDF (usato da menu File > Esporta PDF)."""
        self._export_pdf()

    def _export_pdf(self, path: Optional[Path] = None) -> bool:
        if not isinstance(path, Path):
            path = None
        if path is None:
            if self.file_path:
                default = str(self.file_path.with_suffix(""))
            else:
                default = str(Path.home() / "documento")
            dest, _ = QFileDialog.getSaveFileName(
                self, tr("richtext.export_pdf"), default, "PDF (*.pdf)"
            )
            if not dest:
                return False
            path = Path(dest)
            if path.suffix.lower() != ".pdf":
                path = path.with_suffix(".pdf")
        if not WEBENGINE_OK:
            return False

        # Estrae l'HTML dal documento Jodit e lo carica in una QWebEngineView
        # pulita (senza toolbar Jodit) per generare un PDF fedele al contenuto.
        html = self.get_html()
        clean_html = RichTextIO._full_html(html)

        _path = path
        from core.webengine import safe_webview

        _view = safe_webview()
        if _view is None:
            return False
        loop = QEventLoop()
        _ok = [False]

        def _on_load(success: bool) -> None:
            if not success:
                loop.quit()
                return

            def _on_pdf(pdf_bytes) -> None:
                if pdf_bytes:
                    try:
                        _path.write_bytes(bytes(pdf_bytes))
                        _ok[0] = True
                    except Exception:
                        pass
                loop.quit()

            _view.page().printToPdf(_on_pdf)

        _view.loadFinished.connect(_on_load)
        _view.setHtml(clean_html, QUrl("about:blank"))
        loop.exec()
        _view.deleteLater()
        return _ok[0]

    # ── Helpers interni ───────────────────────────────────────────────────────

    def _request_save(self, path: Path, save_as: bool) -> bool:
        if (
            self._closed
            or self._io_worker is not None
            or self._io_pending
            or not WEBENGINE_OK
            or not self._js_ready
        ):
            return False
        self._io_pending = True
        self._toolbar.setEnabled(False)
        self._status_lbl.setText(tr("richtext.saving", default="Salvataggio in corso..."))

        def _got_html(html: str) -> None:
            if not self._closed and self._io_pending and self._io_worker is None:
                self._io_pending = False
                self._save_change_serial = self._change_serial
                self._io_save_as = save_as
                self._start_io("save", path, html or "")

        self._view.page().runJavaScript("window.getContent()", _got_html)
        return True

    def _start_io(self, operation: str, path: Path, html: str = "") -> bool:
        if self._closed or self._io_worker is not None or self._io_pending:
            return False
        self._io_operation = operation
        self._io_path = path
        self._toolbar.setEnabled(False)
        self._progress_dlg = QProgressDialog(
            tr("richtext.loading", default="Operazione rich text in corso..."),
            tr("button.cancel", default="Annulla"), 0, 0, self,
        )
        self._progress_dlg.setWindowTitle("NotePadPQ")
        self._progress_dlg.setWindowModality(Qt.WindowModality.WindowModal)
        self._progress_dlg.setMinimumDuration(300)
        worker = _RichTextIOWorker(operation, path, html)
        self._io_worker = worker
        self._progress_dlg.canceled.connect(worker.cancel)
        worker.completed.connect(self._on_io_completed)
        worker.finished.connect(self._on_io_thread_finished)
        _keep_thread_alive(worker)
        worker.start()
        return True

    @pyqtSlot(str, str, bool)
    def _on_io_completed(self, result: str, error: str, cancelled: bool) -> None:
        operation, path = self._io_operation, self._io_path
        ok = not cancelled and not error
        self._io_ok = ok
        if error:
            key = "richtext.load_error" if operation == "load" else "richtext.save_error"
            QMessageBox.critical(self, "NotePadPQ", tr(key, err=error))
        elif ok and operation == "load" and path is not None:
            self.file_path = path
            self.set_html(result)
            self._set_modified(False)
            self._update_status()
        elif ok and operation == "save" and path is not None:
            if self._io_save_as:
                self.file_path = path
            if self._change_serial == self._save_change_serial:
                self._set_modified(False)
            self._update_status()

    @pyqtSlot()
    def _on_io_thread_finished(self) -> None:
        operation, path, ok = self._io_operation, self._io_path, self._io_ok
        if self._progress_dlg is not None:
            self._progress_dlg.close()
            self._progress_dlg.deleteLater()
            self._progress_dlg = None
        worker = self._io_worker
        self._io_worker = None
        self._io_operation = ""
        self._io_path = None
        self._io_save_as = False
        self._io_ok = False
        self._toolbar.setEnabled(True)
        self._update_status()
        if worker is not None:
            worker.deleteLater()
        if operation == "load":
            self.load_finished.emit(ok)
        else:
            self.save_finished.emit(ok, path)

    def _js_set_content(self, html: str) -> None:
        if self._view is None:
            return
        escaped = json.dumps(html)
        self._view.page().runJavaScript(f"window.setContent({escaped})")

    def _set_modified(self, mod: bool) -> None:
        if mod != self._modified:
            self._modified = mod
            self.modified_changed.emit(mod)
            self._mod_lbl.setText(("● " + tr("label.modified")) if mod else "")

    def _update_status(self) -> None:
        if self.file_path:
            self._status_lbl.setText(str(self.file_path))

    def _refresh_word_count(self) -> None:
        if not WEBENGINE_OK or not self._js_ready:
            return

        def cb(html: str):
            if not html:
                return
            import re

            text = re.sub(r"<[^>]+>", " ", html)
            words = len(text.split())
            self._wc_label.setText(tr("richtext.word_count", count=words))

        self._view.page().runJavaScript("window.getContent()", cb)

    def _open_as_text(self) -> None:
        """Apre il contenuto HTML in un tab testo senza chiudere il richeditor."""
        name = (self.file_path.stem + ".html") if self.file_path else "documento.html"
        if not WEBENGINE_OK or not self._js_ready:
            self.convert_to_text.emit("", name)
            return
        # Lettura asincrona: emette il segnale dal callback JS senza loop.exec()
        _name = name

        def _cb(html: str):
            self.convert_to_text.emit(html or "", _name)

        self._view.page().runJavaScript("window.getContent()", _cb)

    # ── Cleanup ───────────────────────────────────────────────────────────────

    def cleanup(self) -> None:
        """Ferma callback e operazioni pendenti prima di rimuovere il tab."""
        if self._closed:
            return
        self._closed = True
        self._io_pending = False
        if self._io_worker is not None:
            self._io_worker.cancel()
        if self._progress_dlg is not None:
            self._progress_dlg.close()
        timer = getattr(self, "_wc_timer", None)
        if timer is not None:
            timer.stop()
        if self._tmpdir and self._tmpdir.exists():
            shutil.rmtree(self._tmpdir, ignore_errors=True)
        self._tmpdir = None

    def closeEvent(self, event):
        self.cleanup()
        super().closeEvent(event)

    def __del__(self):
        tmpdir = getattr(self, "_tmpdir", None)
        if tmpdir and tmpdir.exists():
            shutil.rmtree(tmpdir, ignore_errors=True)
